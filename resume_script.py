from docx import Document
from docx.shared import Pt
from ollama import chat
from docx2pdf import convert

import shutil
import os
import re
import time
import json
import hashlib


# ================== CONFIGURATION ==================

MODEL = "llama3.1:8b"

DELAY_BETWEEN_CALLS = 0.7

CACHE_FILE = "resume_cache.json"

# ===================================================



def check_and_pull_model(model_name):

    print(f"Checking model: {model_name}")

    try:

        chat(
            model=model_name,
            messages=[
                {
                    "role":"user",
                    "content":"hi"
                }
            ],
            options={
                "num_predict":5
            }
        )

        print("Model ready")


    except Exception as e:

        if "not found" in str(e).lower():

            os.system(
                f"ollama pull {model_name}"
            )

        else:

            print(e)




def load_cache():

    if os.path.exists(CACHE_FILE):

        try:

            with open(
                CACHE_FILE,
                "r",
                encoding="utf-8"
            ) as f:

                return json.load(f)

        except:

            return {}

    return {}




def save_cache(cache):

    with open(
        CACHE_FILE,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            cache,
            f,
            indent=2
        )




def get_cache_key(text, keywords):

    return hashlib.md5(
        (text+"|"+keywords).encode()
    ).hexdigest()




def extract_ats_keywords(job_description):

    prompt = f"""

Extract the 20 most important ATS keywords,
hard skills, and phrases.

Return ONLY comma separated values.

JOB DESCRIPTION:

{job_description}

"""


    response = chat(

        model=MODEL,

        messages=[
            {
                "role":"user",
                "content":prompt
            }
        ],

        options={
            "temperature":0.1,
            "num_predict":300
        }

    )


    return response["message"]["content"].strip()




def is_likely_bullet(text):

    if len(text)<25:
        return False


    bullets=[
        '•','·','-',
        '▪','■',
        '◆','*'
    ]


    if any(
        text.startswith(x)
        for x in bullets
    ):
        return True



    verbs={
        "led",
        "managed",
        "developed",
        "created",
        "implemented",
        "built",
        "designed",
        "optimized",
        "analyzed",
        "engineered",
        "automated",
        "improved"
    }


    return (
        text.split()[0].lower()
        in verbs
    )




def process_paragraph_for_extraction(
    text,
    para,
    section,
    role,
    bullets
):

    if not text:
        return


    if (
        "skills" in section.lower()
        and
        ("," in text or ":" in text)
        and
        not is_likely_bullet(text)
    ):

        bullets.append({

            "original":text,
            "clean":text,
            "section":section,
            "role":"Technical Skills",
            "is_skills":True

        })

        return



    if is_likely_bullet(text):

        clean = re.sub(
            r'^[•·\-▪■◆*–—\s]+',
            '',
            text
        ).strip()


        bullets.append({

            "original":text,
            "clean":clean,
            "section":section,
            "role":role,
            "is_skills":False

        })




def extract_bullets_with_context(doc):

    bullets=[]

    section=""
    role=""


    for para in doc.paragraphs:


        text=para.text.strip()


        if not text:
            continue



        if len(text)<70 and any(

            x in text.lower()

            for x in [
                "experience",
                "projects",
                "skills",
                "education"
            ]

        ):

            section=text
            continue



        process_paragraph_for_extraction(
            text,
            para,
            section,
            role,
            bullets
        )



    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for p in cell.paragraphs:

                    process_paragraph_for_extraction(
                        p.text.strip(),
                        p,
                        section,
                        role,
                        bullets
                    )


    return bullets




def rewrite_single_bullet(
    bullet,
    job_keywords,
    cache
):

    cache_key = get_cache_key(
        bullet['clean'],
        job_keywords
    )

    if cache_key in cache:
        print("   [CACHE HIT]")
        return cache[cache_key]


    # ================= SKILLS OPTIMIZATION =================

    if bullet.get('is_skills', False):

        prompt = f"""
You are an expert resume writer and ATS optimization specialist.

ATS KEYWORDS: {job_keywords}
ORIGINAL SKILLS: {bullet['clean']}

TASK:
Rewrite this technical skills list to optimize for the ATS keywords.

RULES:
- Reorder the list so that skills matching the ATS keywords appear FIRST.
- Rename existing skills to perfectly match the ATS keyword phrasing if they are synonymous (e.g., change "Node" to "Node.js", or "AWS" to "Amazon Web Services").
- NEVER invent or add new skills that are not present in the original list. Only reorder and reformat existing skills.
- Maintain the original structural formatting (e.g., if categorized like "Languages: Python, Java", keep the categories and colons intact).

OUTPUT FORMATTING:
- Return ONLY the rewritten skills text.
- Do NOT include explanations, quotes, or JSON.
"""


    # ================= BULLET OPTIMIZATION =================

    else:

        prompt = f"""
You are an expert resume writer and ATS optimization specialist.

SECTION CONTEXT: {bullet['section']}
ROLE CONTEXT: {bullet['role']}
ATS KEYWORDS: {job_keywords}

ORIGINAL BULLET:
{bullet['clean']}

TASK:
Rewrite the bullet to naturally integrate relevant ATS keywords while preserving the exact core accomplishment.

RULES:
- Maintain the exact same metrics, business impact, and original core meaning.
- Preserve all existing technologies, numbers, and quantities.
- Never invent new skills, outcomes, or responsibilities.
- Only use ATS keywords if they accurately fit the context of the original bullet.
- Start with a strong, past-tense action verb (unless it's a current role, then use present tense).
- Ensure perfect grammar and avoid AI-sounding buzzwords (e.g., "leveraging", "spearheaded").
- Keep the length within ±20% of the original.

OUTPUT FORMATTING:
- Return ONLY the rewritten text.
- Do NOT include bullet point characters (like -, •, or *) at the beginning. Word handles formatting natively.
- Do NOT include explanations, quotes, or JSON.
"""


    try:

        response = chat(

            model=MODEL,

            messages=[

                {
                    "role": "system",
                    "content":
                    "You are a precise resume optimization expert. You output only finalized text."
                },

                {
                    "role": "user",
                    "content": prompt
                }

            ],

            options={
                "temperature":0.15,
                "num_predict":150
            }

        )


        new_bullet = (
            response["message"]["content"]
            .strip()
        )


        # remove accidental bullet symbols
        if not bullet.get('is_skills', False):

            new_bullet = re.sub(
                r'^[-•·▪■◆*–—\s]+',
                '',
                new_bullet
            ).strip()



        if new_bullet:

            cache[cache_key] = new_bullet

            return new_bullet


    except Exception as e:

        print(
            f"⚠️ Rewrite error: {e}"
        )


    return bullet['clean']



def replace_text_and_shrink(
    paragraph,
    old,
    new
):

    if paragraph.text.strip()==old:


        style=paragraph.style


        paragraph.text=new

        paragraph.style=style


        for run in paragraph.runs:

            run.font.size=Pt(9.5)


        return True


    return False




def tailor_resume(
    resume_path,
    job_description,
    output_docx,
    output_pdf
):


    print("Starting resume tailoring")


    check_and_pull_model(MODEL)


    cache=load_cache()



    doc=Document(resume_path)



    keywords=extract_ats_keywords(
        job_description
    )


    bullets=extract_bullets_with_context(doc)



    changes={}



    for bullet in bullets:


        rewritten=rewrite_single_bullet(
            bullet,
            keywords,
            cache
        )


        if rewritten:

            changes[
                bullet["original"]
            ] = rewritten



        time.sleep(
            DELAY_BETWEEN_CALLS
        )



    save_cache(cache)



    shutil.copy2(
        resume_path,
        output_docx
    )



    new_doc=Document(output_docx)



    replaced=0



    for para in new_doc.paragraphs:


        if para.text.strip() in changes:


            if replace_text_and_shrink(
                para,
                para.text.strip(),
                changes[para.text.strip()]
            ):

                replaced+=1




    for table in new_doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:


                    if para.text.strip() in changes:


                        if replace_text_and_shrink(
                            para,
                            para.text.strip(),
                            changes[para.text.strip()]
                        ):

                            replaced+=1




    new_doc.save(output_docx)



    convert(
        output_docx,
        output_pdf
    )



    print(
        f"Finished. Replaced {replaced} items"
    )


    return output_pdf